import openai
import pandas as pd
import os
import time
from docx import Document

openai.api_key = #Open API Key

output_folder = "reports"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

def generate_report(topic):
    prompt = f"""
Write a business review on {topic} in India. Cover:  
1. **Market Overview**: Key products, demand sectors, and supply-demand trends.  
2. **Key Players**: Major suppliers and market share.  
3. **Growth & Challenges**: Market drivers, constraints, and future outlook.  
Add bullet points when it comes to products or sectors like those things. And Add some numbers to make the report much valueable and professional
Keep it under **1500 words** with relevant data.
"""
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional business analyst."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000  # token limit
        )

        return response.choices[0].message.content

    except openai.OpenAIError as e:
        print(f"API Error: {e}")
        return None

def save_report(topic, content):
    if not content:
        print(f"Skipping saving {topic} due to empty content.")
        return

    doc = Document()
    doc.add_heading(f"Business Review: {topic}", level=1)
    doc.add_paragraph(content)

    filename = f"{output_folder}/{topic.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"Saved: {filename}")

file_path = r"D:\White paper reports\topics.xlsx"  
df = pd.read_excel(file_path)

for index, row in df.iterrows():
    topic = row["Topics"]
    print(f"Generating report for: {topic}...")

    try:
        report_content = generate_report(topic)
        if report_content:
            save_report(topic, report_content)
        else:
            print(f"Failed to generate report for: {topic}")

        time.sleep(5)  # delay - rate limit issues

    except Exception as e:
        print(f"Error generating report for {topic}: {e}")